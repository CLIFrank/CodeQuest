"""Build the Google Docs-targeted CodeQuest technical documentation."""

from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLACK = "000000"
MUTED = "555555"
BORDER = "DADCE0"
CODE_FILL = "F1F3F4"
CONTENT_DXA = 9360


def set_run_font(run, name: str = "Arial", size: float = 11, color: str = BLACK, bold=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    return run


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError("Table column widths must total 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_quiet_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_image_alt(inline_shape, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (20, BLACK, 20, 6),
        "Heading 2": (16, BLACK, 18, 6),
        "Heading 3": (14, "434343", 16, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_body(doc: Document, text: str):
    return doc.add_paragraph(text)


def add_labelled(doc: Document, label: str, text: str):
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run(label), bold=True)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.0
    shade_paragraph(paragraph, CODE_FILL)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        set_run_font(paragraph.add_run(line or " "), "Consolas", 9.5, "202124")


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(10)
    set_run_font(paragraph.add_run(text), "Arial", 9, MUTED)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_quiet_borders(table)
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(value), bold=True)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_run_font(paragraph.add_run(value))
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def font(size: int, bold: bool = False):
    path = Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf")
    return ImageFont.truetype(str(path), size)


def build_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 620), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(42, True)
    box_font = font(28, True)
    small_font = font(23)
    draw.text((70, 40), "CodeQuest application architecture", fill="black", font=title_font)

    boxes = {
        "entry": (70, 165, 300, 285, "main.py", "Starts the app"),
        "app": (385, 145, 705, 305, "CodeQuestApp", "Loop and orchestration"),
        "ui": (805, 100, 1115, 220, "Pygame UI", "Screens and editor"),
        "data": (805, 250, 1115, 370, "Curriculum", "Quest definitions"),
        "state": (805, 400, 1115, 520, "Progress", "JSON persistence"),
        "runner": (1215, 175, 1530, 335, "Code runner", "Validation + subprocess"),
    }
    for key, (x1, y1, x2, y2, heading, detail) in boxes.items():
        fill = "#F1F3F4" if key != "app" else "#E8F0FE"
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline="#DADCE0", width=3)
        heading_width = draw.textbbox((0, 0), heading, font=box_font)[2]
        detail_width = draw.textbbox((0, 0), detail, font=small_font)[2]
        draw.text(((x1 + x2 - heading_width) / 2, y1 + 28), heading, fill="black", font=box_font)
        draw.text(((x1 + x2 - detail_width) / 2, y1 + 74), detail, fill="#555555", font=small_font)

    def arrow(start, end):
        draw.line((start, end), fill="#5F6368", width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#5F6368")

    arrow((300, 225), (385, 225))
    arrow((705, 195), (805, 160))
    arrow((705, 225), (805, 310))
    arrow((705, 255), (805, 460))
    arrow((1115, 160), (1215, 225))
    arrow((1115, 310), (1215, 260))
    arrow((1215, 310), (1115, 445))
    draw.text((1185, 370), "Results update progress", fill="#555555", font=small_font)
    image.save(path, dpi=(220, 220))


def build(output: Path, diagram: Path, screenshot: Path | None) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    diagram.parent.mkdir(parents=True, exist_ok=True)
    build_architecture_diagram(diagram)

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "CodeQuest Development Documentation and Technical Analysis"
    props.subject = "Architecture and implementation analysis of the Pygame Python learning application"
    props.keywords = "CodeQuest, Pygame, Python, education, technical documentation"

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    set_run_font(title.add_run("CodeQuest Development Documentation"), "Arial", 26, BLACK, False)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(subtitle.add_run("Technical breakdown and implementation analysis"), "Arial", 14, MUTED, False)
    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(16)
    set_run_font(metadata.add_run("Pygame Python learning application | Documentation date: July 31, 2026"), "Arial", 10, MUTED)

    doc.add_heading("Executive summary", level=1)
    add_body(
        doc,
        "CodeQuest is a data-driven desktop learning environment built with Pygame. It teaches beginner Python through short narrative missions in which a learner edits real Python, executes it in a restricted subprocess, compares the output with a target, earns XP, and unlocks later levels. The same application also tracks personal study tasks and retains learner code between sessions.",
    )
    add_body(
        doc,
        "The implementation deliberately separates curriculum data, learner state, code execution, interface controls, and screen orchestration. This makes the application easier to test and extend than a single monolithic Pygame script. The present build includes eight progressive quests, seven passing automated tests, and a visually verified dashboard-to-completion workflow.",
    )

    if screenshot and screenshot.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(str(screenshot), width=Inches(6.35))
        set_image_alt(shape, "CodeQuest dashboard showing the current mission, quest map, XP, and learner task panel")
        add_caption(doc, "Figure 1. CodeQuest learner dashboard.")

    doc.add_heading("Product interpretation and scope", level=1)
    add_body(doc, "The initial concept was translated into five concrete learner capabilities:")
    add_numbered(
        doc,
        [
            "Select and unlock Python game levels.",
            "Write Python inside the application.",
            "Run code and see output or readable errors.",
            "Submit solutions, earn XP, and progress through a curriculum.",
            "Create and complete personal learning tasks.",
        ],
    )
    add_body(
        doc,
        "The workspace was initially empty, so the project was designed and implemented from scratch. The first version intentionally remains a local desktop application: it has no server, authentication system, external database, or remote execution service.",
    )

    doc.add_heading("System architecture", level=1)
    add_body(
        doc,
        "CodeQuest follows a small component architecture. CodeQuestApp owns the real-time application state and coordinates specialized modules instead of implementing curriculum, persistence, execution, and editing logic itself.",
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(diagram), width=Inches(6.4))
    set_image_alt(shape, "Architecture diagram linking main.py, CodeQuestApp, Pygame UI, curriculum, progress, and the code runner")
    add_caption(doc, "Figure 2. High-level application architecture and data flow.")

    add_table(
        doc,
        ["Component", "Responsibility"],
        [
            ["main.py", "Minimal application entry point."],
            ["codequest/app.py", "Game loop, navigation, screen rendering, assessment flow, and celebrations."],
            ["codequest/widgets.py", "Buttons, task input, and the custom multiline code editor."],
            ["codequest/theme.py", "Dimensions, colors, fonts, rounded panels, text drawing, and wrapping."],
            ["codequest/curriculum.py", "Quest model, eight built-in missions, and output comparison."],
            ["codequest/runner.py", "AST validation, subprocess orchestration, timeout, and result parsing."],
            ["codequest/runner_worker.py", "Restricted built-ins, learner-code execution, and stdout capture."],
            ["codequest/progress.py", "XP, sequential unlocks, learner tasks, saved code, and JSON persistence."],
            ["tests/", "Curriculum, runner, progression, and persistence verification."],
        ],
        [2700, 6660],
    )

    doc.add_heading("Startup and application lifecycle", level=1)
    add_body(
        doc,
        "main.py imports CodeQuestApp, constructs it, and invokes run(). During initialization, Pygame and its font subsystem are initialized, clipboard support is attempted, and a resizable 1280 by 760 window is created. The application then loads local learner progress, selects the next available quest, restores saved code or starter code, and initializes the console, modal, hint, toast, and animation state.",
    )
    add_labelled(doc, "Primary reference: ", "codequest/app.py:53")
    add_body(doc, "The main loop runs at a target of 60 frames per second and uses the sequence:")
    add_code(doc, "collect events -> update animation state -> draw frame -> present frame -> cap frame rate")
    add_body(
        doc,
        "The optional max_frames argument exists so the interface can be exercised in automated headless smoke tests. On normal shutdown, active editor text and the full learner profile are saved before Pygame closes.",
    )

    doc.add_heading("Event handling and navigation", level=1)
    add_body(doc, "Input is routed centrally and in a deliberate priority order:")
    add_numbered(
        doc,
        [
            "Window close and resize events.",
            "Success-modal controls.",
            "New-task modal controls.",
            "Visible button callbacks.",
            "Mission keyboard shortcuts and code-editor input.",
        ],
    )
    add_body(
        doc,
        "This ordering prevents an event intended for a modal from activating a control underneath it. Ctrl+Enter checks a solution, Escape returns to the dashboard, mouse clicks activate missions and controls, and the task dialog supports Enter and Escape. Resizing is constrained to at least 1050 by 680 pixels so the principal panels remain usable.",
    )
    add_labelled(doc, "Primary reference: ", "codequest/app.py:102")

    doc.add_heading("Rendering and interface composition", level=1)
    add_body(
        doc,
        "Each frame draws the background, sidebar, active screen, confetti, temporary notifications, and any active modal. Buttons are reconstructed during the draw pass so their hit boxes always correspond to the current responsive layout.",
    )
    doc.add_heading("Dashboard", level=2)
    add_bullets(
        doc,
        [
            "Displays XP, learner level, streak, and total journey progress.",
            "Highlights the current mission and its concept and reward.",
            "Generates the quest map directly from curriculum data.",
            "Uses distinct locked, available, hovered, and completed states.",
            "Provides a personal task panel with create and completion actions.",
        ],
    )
    doc.add_heading("Mission workspace", level=2)
    add_bullets(
        doc,
        [
            "Shows narrative context, objective, expected output, reward, and optional hint.",
            "Places the editable Python program beside the instructions.",
            "Separates running from graded submission through Run and Check actions.",
            "Shows normal output, validation errors, Python exceptions, or comparison feedback in the console.",
            "Provides Reset, Back, and keyboard shortcut controls.",
        ],
    )
    doc.add_heading("Completion feedback", level=2)
    add_body(
        doc,
        "A successful check persists progress, opens a mission-complete modal, displays total XP and completed-quest count, and launches a lightweight confetti particle system. All decorative visuals are generated with Pygame primitives; the application does not depend on external UI artwork.",
    )

    doc.add_heading("Visual design system", level=1)
    add_body(
        doc,
        "theme.py centralizes the 1280 by 760 base resolution, 224-pixel sidebar, colors, font fallbacks, rounded panels, anchored text, and word wrapping. The interface uses dark navy surfaces, purple primary actions, cyan information highlights, green success states, yellow hints, red errors, and muted secondary text.",
    )
    add_body(
        doc,
        "Segoe UI and Arial are requested for interface text, while Consolas or another monospaced fallback is requested for code. A later visual QA pass replaced unsupported font glyphs with geometric Pygame line drawings so status icons render consistently on Windows.",
    )
    add_labelled(doc, "Primary reference: ", "codequest/theme.py:10")

    doc.add_heading("Custom Python editor", level=1)
    add_body(
        doc,
        "Pygame does not provide a native multiline editor, so CodeQuest implements one in widgets.py. Its core state consists of the complete source string, an integer cursor offset, focus state, vertical scroll offset, preferred vertical-navigation column, and up to 100 undo snapshots.",
    )
    add_bullets(
        doc,
        [
            "Character input through Pygame TEXTINPUT events.",
            "Backspace, Delete, arrow movement, Home, and End.",
            "Four-space Tab insertion and automatic indentation after a colon.",
            "Indentation inheritance on new lines.",
            "Ctrl+Z undo and clipboard paste.",
            "Whole-document copy, mouse cursor placement, and mouse-wheel scrolling.",
            "Line numbers, cursor blinking, and visible-line management.",
        ],
    )
    add_body(
        doc,
        "Syntax colors are calculated with Python's tokenize module. Strings, numbers, comments, operators, keywords, and selected built-ins receive distinct colors. Tokenization and indentation errors are caught because incomplete code is normal while a learner is typing. In that case, the editor continues drawing whatever tokens it can recognize.",
    )
    add_labelled(doc, "Primary references: ", "codequest/widgets.py:116 and codequest/widgets.py:298")
    add_body(
        doc,
        "Current editor limitations are no text selection, no partial cut or copy, no search, no autocomplete, and no language-server integration.",
    )

    doc.add_heading("Curriculum model and learning progression", level=1)
    add_body(
        doc,
        "Each mission is an immutable Quest record with a stable ID, world, number, title, concept, story, objective, starter code, expected output, hint, XP reward, and difficulty. UI screens derive their content from these records rather than maintaining duplicate labels and values.",
    )
    add_table(
        doc,
        ["#", "Concept", "Mission", "Learning objective"],
        [
            ["1", "print()", "Send a signal", "Produce the learner's first exact text output."],
            ["2", "Variables", "Choose a codename", "Store a string and print the stored value."],
            ["3", "Numbers", "Fuel the rocket", "Combine numeric variables with addition."],
            ["4", "Loops", "Launch countdown", "Repeat output with a descending range."],
            ["5", "Conditionals", "Open the airlock", "Evaluate a numeric safety condition."],
            ["6", "Lists", "Pack the cargo", "Construct and iterate through a sequence."],
            ["7", "Functions", "Alien translator", "Define, return from, and call a function."],
            ["8", "Dictionaries", "Decode the star map", "Store and retrieve a keyed value."],
        ],
        [620, 1600, 2500, 4640],
    )
    add_body(
        doc,
        "Starter programs are intentionally incomplete. They provide scaffolding without directly solving the mission. Hints reveal the relevant syntax but remain optional.",
    )
    add_body(
        doc,
        "Output grading removes surrounding blank space and trailing whitespace on individual lines, while preserving line order, capitalization, and meaningful internal spaces. An extra final newline is accepted, but hello does not match Hello.",
    )
    add_labelled(doc, "Primary references: ", "codequest/curriculum.py:9 and codequest/curriculum.py:145")

    doc.add_heading("Learner-code execution", level=1)
    add_body(
        doc,
        "Execution is intentionally separated from the Pygame process. The application first performs static checks, then launches a short-lived isolated Python subprocess and exchanges plain text and JSON with it.",
    )
    doc.add_heading("Static AST validation", level=2)
    add_body(doc, "Before launching Python, validate_code parses the learner program into an abstract syntax tree and rejects:")
    add_bullets(
        doc,
        [
            "Programs longer than 10,000 characters.",
            "Import and from-import statements.",
            "Double-underscore names and double-underscore attribute access.",
            "Calls to open, eval, exec, compile, input, globals, locals, vars, breakpoint, and __import__.",
        ],
    )
    add_body(
        doc,
        "Syntax errors are allowed to reach the worker so Python can return the familiar exception message rather than a generic validator message.",
    )
    doc.add_heading("Subprocess worker", level=2)
    add_numbered(
        doc,
        [
            "runner.py launches the current interpreter with isolated mode enabled through -I.",
            "Learner code is transmitted through standard input rather than a temporary source file.",
            "runner_worker.py compiles and executes the code with a deliberately small built-in dictionary.",
            "Standard output is captured in memory.",
            "Success, output, and exception information are serialized as JSON.",
            "The parent terminates the run after two seconds and reports a likely endless loop.",
        ],
    )
    add_body(
        doc,
        "The exposed built-ins cover print, range, basic types and collections, enumerate, zip, sum, min, max, abs, round, and sorted. On Windows the child process is configured not to open a separate console window.",
    )
    add_labelled(doc, "Primary references: ", "codequest/runner.py:34 and codequest/runner_worker.py:17")

    doc.add_heading("Security analysis", level=2)
    add_body(
        doc,
        "The runner prevents common classroom accidents and makes exercises predictable, but it is an educational guardrail rather than a hardened security boundary. AST filters and reduced built-ins should not be treated as sufficient protection against determined hostile users. A shared or internet-facing deployment should execute code in an operating-system sandbox or disposable container with explicit CPU, memory, filesystem, process, and network limits.",
    )

    doc.add_heading("Run and submission flow", level=1)
    doc.add_heading("Run", level=2)
    add_numbered(
        doc,
        [
            "Copy the current editor source into the in-memory learner profile.",
            "Execute it through the subprocess runner.",
            "Display captured output if it succeeds.",
            "Display validation or Python exception details if it fails.",
        ],
    )
    add_body(doc, "Run does not alter XP, completed missions, or unlock state.")
    doc.add_heading("Check", level=2)
    add_numbered(
        doc,
        [
            "Execute the program through the same runner.",
            "Stop if validation or execution fails.",
            "Normalize and compare actual output with the quest target.",
            "Show actual and expected output when they differ.",
            "Mark the mission complete when they match.",
            "Award XP only for the first successful completion.",
            "Persist progress and display the completion celebration.",
        ],
    )
    add_body(
        doc,
        "The current grader is output-based. This keeps early exercises understandable, but it means a learner can pass by hard-coding the expected output. A stronger assessment system would support hidden test cases, function-level checks, or quest-specific graders.",
    )
    add_labelled(doc, "Primary reference: ", "codequest/app.py:634")

    doc.add_heading("Progression, tasks, and persistence", level=1)
    add_body(
        doc,
        "Progress contains completed quest IDs, XP, streak, personal tasks, and a source-code dictionary keyed by quest ID. The first quest is always unlocked; each later quest requires completion of the immediately preceding one. The complete method checks existing IDs before awarding XP, which prevents repeated farming of the same reward.",
    )
    add_body(
        doc,
        "Personal tasks receive monotonically increasing IDs, are trimmed to 80 characters, and contain a done flag. They can be created and toggled from the dashboard.",
    )
    add_body(doc, "The default save location is:")
    add_code(doc, "~/.codequest/progress.json")
    add_body(
        doc,
        "Saving writes JSON to a temporary sibling file and then replaces the real save file, reducing the chance of partial data after interruption. Loading validates quest IDs and value types. Missing, malformed, or incompatible data produces a fresh learner profile rather than crashing the application.",
    )
    add_body(
        doc,
        "ProgressStore also accepts a custom path. Tests use this dependency-injection point to create isolated temporary profiles without modifying real learner progress.",
    )
    add_labelled(doc, "Primary references: ", "codequest/progress.py:21 and codequest/progress.py:49")

    doc.add_heading("Verification and quality assurance", level=1)
    add_body(doc, "Seven automated tests currently pass. They verify:")
    add_bullets(
        doc,
        [
            "Every built-in quest has a working reference solution.",
            "Each reference solution produces the required output.",
            "Output normalization removes only intended trailing whitespace.",
            "Imports and direct file access are rejected.",
            "Python exceptions are returned in learner-readable form.",
            "Quests unlock in order and duplicate completion does not duplicate XP.",
            "Tasks, XP, and source code survive a save-and-load round trip.",
            "Invalid JSON produces fresh progress safely.",
        ],
    )
    add_body(doc, "The repeatable verification commands are:")
    add_code(doc, "python -m compileall -q .\npython -m unittest discover -s tests -v")
    add_body(
        doc,
        "A separate headless Pygame smoke test rendered the dashboard and mission screens, executed a correct answer, confirmed progress mutation, and rendered the success modal. The generated screenshots were manually inspected for clipping, overlap, icon rendering, spacing, and state consistency.",
    )

    doc.add_heading("Current limitations and risks", level=1)
    add_bullets(
        doc,
        [
            "Output-only grading can be bypassed with hard-coded output.",
            "The execution call runs synchronously and may pause the interface for up to two seconds.",
            "The restricted runner is not sufficient for hostile remote code.",
            "The application supports one local learner profile and no authentication.",
            "Streak is stored but not yet calculated from calendar activity.",
            "Curriculum authoring currently requires editing Python data.",
            "The editor does not offer selection, autocomplete, search, or debugging.",
            "Progress is local and has no cloud synchronization or backup.",
            "An abrupt process crash can lose editor changes made since the most recent disk save.",
        ],
    )

    doc.add_heading("Extension guide", level=1)
    doc.add_heading("Adding a mission", level=2)
    add_numbered(
        doc,
        [
            "Add a Quest record to the QUESTS tuple in codequest/curriculum.py.",
            "Assign a unique stable ID, starter program, expected output, hint, and XP reward.",
            "Add a reference solution to tests/test_curriculum.py.",
            "Run compilation, automated tests, and the UI smoke test.",
        ],
    )
    add_body(
        doc,
        "The quest map, progress denominator, sequential unlock rules, mission panel, hint, code persistence, and XP award all consume curriculum data automatically.",
    )
    doc.add_heading("Adding Python capabilities", level=2)
    add_body(
        doc,
        "If a new exercise requires another built-in, it must be added deliberately to SAFE_BUILTINS in runner_worker.py and covered by a test. Import support should not be enabled casually because it substantially changes the execution threat model.",
    )
    doc.add_heading("Recommended production roadmap", level=2)
    add_numbered(
        doc,
        [
            "Move learner execution into a hardened container or operating-system sandbox.",
            "Introduce asynchronous execution so timeouts never freeze the interface.",
            "Add quest-specific hidden tests and semantic grading.",
            "Create learner profiles, teacher controls, and curriculum-authoring tools.",
            "Add structured analytics, activity-derived streaks, and cloud synchronization.",
            "Expand editor accessibility, selection, autocomplete, and debugging support.",
        ],
    )

    doc.add_heading("Running the application", level=1)
    add_body(doc, "Python 3.10 or later is recommended. Install dependencies and start the application with:")
    add_code(doc, "python -m pip install -r requirements.txt\npython main.py")
    add_body(doc, "Key controls:")
    add_bullets(
        doc,
        [
            "Click an unlocked mission card to open it.",
            "Use Run to preview output and Check to submit.",
            "Press Ctrl+Enter to check from the keyboard.",
            "Press Tab for four spaces and Ctrl+Z to undo.",
            "Press Escape to return to the dashboard.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    add_body(
        doc,
        "CodeQuest is a complete local prototype with a coherent learner journey, a reusable curriculum model, a functional in-app editor, guarded execution, persistent progression, and automated verification. Its architecture is intentionally modest but provides clear seams for stronger grading, secure deployment, teacher tooling, and multi-user services without requiring the interface to be rewritten from the beginning.",
    )

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--diagram", type=Path, required=True)
    parser.add_argument("--screenshot", type=Path)
    args = parser.parse_args()
    build(args.output, args.diagram, args.screenshot)


if __name__ == "__main__":
    main()
